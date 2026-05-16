from fastapi import FastAPI, Request, Response
import httpx, os, fitz
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

app = FastAPI()

# --- CONFIG ---
ACCESS_TOKEN = "EAAavL5WiIDcBRVamTWxrGHxJTEaLDK6KsWF2l9aTcT6vZAquKI8UAmkhZBhulJz8OPzHrkFKIU2S7x9ZC7RnXamywUrJqj64mkcLlZBik4cMuYxFk10yjaZBOIo8RzV2gmW65Eh1fIr1Wsdul5pq7qZA9j3F1JpsCWcHC7RuEsv90YK7GGRnNRazhEXpqN3ovYJ3pFs4hkqrKD4VljvOLRIowM5NsVyZAUxZCsP552wOnOuqZCImcimMdz2DQJnGMDhNnAYFASZBfxWDbLTUHOfIlvDRKSvS0KyTrWGVZB3kAZDZD"
PHONE_NUMBER_ID = "1103451669516095"
VERIFY_TOKEN = "MY_SECRET_TOKEN_123"
groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

db = {}

# --- THE DESIGNER ENGINE (Modern Layout) ---
def build_premium_mzansi_doc(cv_raw, letter_raw):
    doc = Document()
    # Setting Narrow Margins for a Modern Look
    sections = doc.sections
    for section in sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)

    navy = RGBColor(0, 51, 102)

    # 1. HEADER (Centered, Bold, Modern)
    name = cv_raw.split('\n')[0].replace('[NAME]', '').strip()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(name)
    run.bold, run.font.size, run.font.color.rgb = True, Pt(24), navy
    
    # 2. TWO-COLUMN TABLE (Sidebar and Body)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_left = table.columns[0]
    col_right = table.columns[1]
    col_left.width = Inches(2.0)
    col_right.width = Inches(5.0)
    
    cells = table.rows[0].cells
    
    # Left Sidebar (Contact & Skills)
    cells[0].paragraphs[0].text = "CONTACT & SKILLS"
    cells[0].paragraphs[0].runs[0].bold = True

    # Right Body (Experience & Education)
    body = cells[1]
    for line in cv_raw.split('\n')[1:]:
        line = line.strip()
        if not line: continue
        if '[SECTION]' in line:
            p = body.add_paragraph(line.replace('[SECTION]', ''))
            run = p.add_run()
            run.bold, run.font.size, run.font.color.rgb = True, Pt(14), navy
        elif '[JOB]' in line:
            p = body.add_paragraph(line.replace('[JOB]', ''))
            p.runs[0].bold = True
        else:
            body.add_paragraph(line, style='List Bullet' if line.startswith('-') else None)

    # 3. POPIA FOOTER
    doc.add_paragraph("\n" + "_"*40)
    footer = doc.add_paragraph("Legal Notice: This document is POPIA compliant. Personal information is provided for the sole purpose of recruitment.")
    footer.runs[0].font.size, footer.runs[0].italic = Pt(8), True

    # 4. COVER LETTER (Page 2)
    if letter_raw:
        doc.add_page_break()
        cl = doc.add_heading("PROFESSIONAL COVER LETTER", level=1)
        cl.runs[0].font.color.rgb = navy
        doc.add_paragraph(letter_raw)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

async def fetch_jobs(cv_text):
    # Improved extraction for Finance/Admin/Payroll
    p = f"Based on this CV, what is a general job search keyword (e.g. 'Payroll Administrator')? Answer 1-2 words only: {cv_text[:500]}"
    res = groq_client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role":"user","content":p}])
    keyword = res.choices[0].message.content.strip()
    
    aid, akey = os.environ.get("ADZUNA_APP_ID"), os.environ.get("ADZUNA_APP_KEY")
    url = f"https://api.adzuna.com/v1/api/jobs/za/search/1?app_id={aid}&app_key={akey}&results_per_page=5&what={keyword}"
    async with httpx.AsyncClient() as client:
        r = await client.get(url)
        jobs = r.json().get("results", [])
        return "\n\n".join([f"📍 *{j['title']}*\n🏢 {j['company']['display_name']}\n🔗 {j['redirect_url']}" for j in jobs])

@app.post("/webhook")
async def receive(request: Request):
    try:
        data = await request.json()
        val = data.get("entry", [{}])[0].get("changes", [{}])[0].get("value", {})
        if "messages" not in val: return {"status": "ok"}
        msg = val["messages"][0]
        phone = msg["from"]
        if phone not in db: db[phone] = {"cv": ""}
        u = db[phone]
        headers = {"Authorization": f"Bearer {ACCESS_TOKEN}"}

        if msg.get("type") == "interactive":
            bid = msg["interactive"]["button_reply"]["id"]
            if bid == "GET_PRO":
                await httpx.AsyncClient().post(f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/messages", headers=headers, json={"messaging_product": "whatsapp", "to": phone, "type": "text", "text": {"body": "🛠️ Architecting your Premium CV & Cover Letter... ⏳"}})
                
                # HIGH IMPACT PROMPTS
                cv_p = f"Act as an Executive CV Consultant. Rewrite this CV using [NAME], [SECTION], and [JOB] markers. Transform boring tasks into high-impact achievements (STAR Method). TEXT: {u['cv'][:2500]}"
                cv_res = groq_client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role":"user","content":cv_p}])
                
                let_p = f"Write a world-class, persuasive South African cover letter for a Senior role based on: {u['cv'][:1000]}"
                let_res = groq_client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role":"user","content":let_p}])
                
                docx = build_premium_mzansi_doc(cv_res.choices[0].message.content, let_res.choices[0].message.content)
                
                async with httpx.AsyncClient() as client:
                    # Upload and Send Document
                    up = await client.post(f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/media", headers=headers, files={"file": ("Mzansi_Executive_Bundle.docx", docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, data={"messaging_product": "whatsapp", "type": "document"})
                    mid = up.json().get("id")
                    await client.post(f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/messages", headers=headers, json={"messaging_product": "whatsapp", "to": phone, "type": "document", "document": {"id": mid, "filename": "Mzansi_Executive_Bundle.docx"}})
                    
                    # Jobs
                    jobs_list = await fetch_jobs(u["cv"])
                    await client.post(f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/messages", headers=headers, json={"messaging_product": "whatsapp", "to": phone, "type": "text", "text": {"body": f"🔥 *ACTIVE OPENINGS FOR YOUR PROFILE:*\n\n{jobs_list}"}})

        elif msg.get("type") == "document":
            async with httpx.AsyncClient() as client:
                r = await client.get(f"https://graph.facebook.com/v18.0/{msg['document']['id']}", headers=headers)
                f_r = await client.get(r.json().get("url"), headers=headers)
                with fitz.open(stream=f_r.content, filetype="pdf") as doc:
                    u["cv"] = "".join([page.get_text() for page in doc])
            
            btn = {"messaging_product": "whatsapp", "to": phone, "type": "interactive", "interactive": {"type": "button", "body": {"text": "✅ CV Read. Ready for your Executive Bundle?"}, "action": {"buttons": [{"type": "reply", "reply": {"id": "GET_PRO", "title": "📥 Download Bundle"}}]}}}
            await httpx.AsyncClient().post(f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/messages", headers=headers, json=btn)
        
        else:
            await send_text(phone, "Welcome! Upload your PDF CV to generate your Premium Mzansi Bundle.")

    except Exception as e: print(f"ERROR: {e}")
    return {"status": "success"}

@app.get("/webhook")
async def verify(request: Request):
    return Response(content=request.query_params.get("hub.challenge"), status_code=200)

async def send_text(to, text):
    url = f"https://graph.facebook.com/v18.0/{PHONE_NUMBER_ID}/messages"
    headers = {"Authorization": f"Bearer {ACCESS_TOKEN}", "Content-Type": "application/json"}
    await httpx.AsyncClient().post(url, headers=headers, json={"messaging_product": "whatsapp", "to": to, "type": "text", "text": {"body": text}})
